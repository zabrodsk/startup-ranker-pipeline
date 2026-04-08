"""Extract text from DOCX files with section-based chunking."""

from pathlib import Path


def extract_docx(path: str | Path) -> list[dict]:
    """Extract text from a DOCX file, returning one item per logical section.

    Headings create section boundaries. If no headings are found, all content
    is returned as a single item. Tables are included inline as pipe-separated rows.

    Args:
        path: Path to the .docx file.

    Returns:
        List of dicts with keys: text, page_or_slide (section number), source_file.
    """
    from docx import Document  # python-docx

    path = Path(path)
    doc = Document(str(path))
    results: list[dict] = []

    current_texts: list[str] = []
    section_num = 1

    def _flush() -> None:
        nonlocal section_num
        combined = "\n".join(current_texts).strip()
        if combined:
            results.append({
                "text": combined,
                "page_or_slide": section_num,
                "source_file": path.name,
            })
            section_num += 1
        current_texts.clear()

    # Walk paragraphs — flush on heading styles to create section boundaries.
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style.name.startswith("Heading") and current_texts:
            _flush()
        current_texts.append(text)

    # Append table content (tables are not interleaved with paragraphs via
    # the high-level API, but they contribute searchable evidence).
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                current_texts.append(row_text)

    _flush()
    return results
